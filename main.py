# main.py

import os
import streamlit as st
import logging
from dotenv import load_dotenv

# Load environment variables first, before any other code
load_dotenv()

# Import configuration defaults (after loading .env to prioritize environment variables)
from config import ENV_DEFAULTS, DEFAULT_CONFIG

# Configure logging based on configuration
log_level = os.environ.get('LOGLEVEL', DEFAULT_CONFIG['LOGLEVEL']).upper()
logging.basicConfig(
    level=getattr(logging, log_level),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        # Only log to console if level is INFO or higher
        logging.StreamHandler() if log_level != 'WARNING' else logging.NullHandler()
    ]
)

# Configure app
st.set_page_config(page_title="Translator & Readability", layout="wide")

# Check for missing environment variables and use defaults from config
for var, default in ENV_DEFAULTS.items():
    if var not in os.environ:
        logging.debug(f"Environment variable {var} not found, using default: {default}")
        os.environ[var] = default

# Model configuration from default config (no hardcoded defaults)
MODEL_CONFIG = {
    "gemma_gpu_layers": DEFAULT_CONFIG["GEMMA_GPU_LAYERS"],
    "gemma_context_size": DEFAULT_CONFIG["GEMMA_CONTEXT_SIZE"],
    "max_parallel_models": DEFAULT_CONFIG["MAX_PARALLEL_MODELS"],
    "session_timeout": DEFAULT_CONFIG["SESSION_TIMEOUT"],
    "allow_gpu": DEFAULT_CONFIG["ALLOW_GPU"]
}

# Initialize model semaphore for limiting concurrent model usage
import threading
model_semaphore = threading.Semaphore(MODEL_CONFIG["max_parallel_models"])

import tempfile
import io
from docx import Document
import uuid
import traceback

from models.nltk_resources import setup_nltk
from utils.file_readers import read_file
from utils.text_processing import detect_language
from utils.readability_indices import (
    flesch_reading_ease,
    flesch_kincaid_grade_level,
    gunning_fog_index,
    smog_index
)
from utils.formatting import color_code_index
from utils.tilmash_translation import tilmash_translate, display_tilmash_streaming_translation
from utils.gemma_translation import gemma_translate, display_streaming_translation

# Initialize session state for user identification
if 'session_id' not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
    
if 'translation_lock' not in st.session_state:
    st.session_state.translation_lock = False

def handle_translation():
    st.header("Перевод (Kazakh, Russian, English)")
    
    # Show session ID in sidebar for debugging
    with st.sidebar.expander("Session Info", expanded=False):
        st.write(f"Session ID: {st.session_state.session_id}")
        
        # Add GPU usage option (admin only)
        if MODEL_CONFIG["allow_gpu"]:
            st.session_state.use_gpu = st.checkbox("Use GPU (faster)", value=False)
        else:
            st.session_state.use_gpu = False
            st.write("GPU mode disabled by configuration")
    
    translate_input_method = st.radio("Способ ввода текста:", ["Загрузить файл", "Вставить текст"])
    input_text = ""

    if translate_input_method == "Загрузить файл":
        uploaded_file = st.file_uploader("Выберите файл (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"])
        if uploaded_file is not None:
            suffix = os.path.splitext(uploaded_file.name)[1]
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                temp_file_path = tmp_file.name
            input_text = read_file(temp_file_path)
            os.remove(temp_file_path)
            st.write("**Содержимое файла:**")
            st.write(input_text)
    else:
        input_text = st.text_area("Вставьте ваш текст здесь", height=200)

    if input_text:
        auto_detect = st.checkbox("Автоматически определить язык", value=True)
        src_lang = None
        if auto_detect:
            detected_lang = detect_language(input_text)
            if detected_lang in ['ru','en','kk']:
                st.info(f"Определён язык: {detected_lang}")
                src_lang = detected_lang
            else:
                st.warning("Не удалось определить язык. Выберите вручную.")
                src_lang = st.selectbox("Язык текста", ["ru", "en", "kk"])
        else:
            src_lang = st.selectbox("Язык текста", ["ru", "en", "kk"])

        if src_lang == "ru":
            tgt_options = ["en","kk"]
        elif src_lang == "en":
            tgt_options = ["ru","kk"]
        else:
            tgt_options = ["ru","en"]

        tgt_lang = st.selectbox("Перевод на:", tgt_options)
        
        # Select translation model
        model_option = st.radio("Выберите модель перевода:", ["Gemma 3", "Tilmash"])

        if st.button("Перевести"):
            # Prevent multiple concurrent translations from same session
            if st.session_state.translation_lock:
                st.warning("Перевод уже выполняется. Пожалуйста, дождитесь завершения.")
                return
                
            # Set translation lock
            st.session_state.translation_lock = True
            
            try:
                # Use the model semaphore to limit concurrent model access
                acquired = model_semaphore.acquire(blocking=False)
                if not acquired:
                    st.warning("Максимальное количество параллельных моделей достигнуто. Пожалуйста, попробуйте позже.")
                    st.session_state.translation_lock = False
                    return
                    
                try:
                    if model_option == "Tilmash":
                        st.subheader("Результат перевода:")
                        # Get the approximate size of the text to determine if chunking is needed
                        approx_text_size = len(input_text) / 4  # rough approximation (4 chars ≈ 1 token)
                        needs_chunking = approx_text_size > 500  # If text is likely over 500 tokens
                        
                        # Display appropriate spinner message
                        spinner_message = "Processing text in chunks..." if needs_chunking else "Processing translation..."
                        
                        try:
                            # Create a dedicated translator instance for this session
                            from utils.tilmash_translation import TilmashTranslator
                            translator = TilmashTranslator()
                            
                            with st.spinner(spinner_message):
                                try:
                                    # Use direct streaming approach with session-specific translator
                                    result = ""
                                    translation_placeholder = st.empty()
                                    
                                    # Stream translation
                                    for chunk in translator.translate_streaming(input_text, src_lang, tgt_lang):
                                        result += chunk
                                        translation_placeholder.markdown(result)
                                        
                                except Exception as e:
                                    st.error(f"Translation error: {str(e)}")
                                    logging.error(f"Tilmash translation error: {traceback.format_exc()}")
                                    result = None
                            
                            if result:
                                # Prepare download capability
                                doc = Document()
                                doc.add_paragraph(result)
                                doc_io = io.BytesIO()
                                doc.save(doc_io)
                                doc_io.seek(0)

                                st.download_button(
                                    label="Скачать переведённый текст (.docx)",
                                    data=doc_io,
                                    file_name="translated_text.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            else:
                                st.warning("Не удалось выполнить перевод.")
                                
                            # Unload Tilmash model after use
                            try:
                                if translator.initialized:
                                    translator.unload_model()
                            except Exception as unload_error:
                                logging.error(f"Error unloading Tilmash model: {str(unload_error)}")
                        except Exception as tilmash_error:
                            st.error(f"Tilmash model error: {str(tilmash_error)}")
                            logging.error(f"Tilmash model error: {traceback.format_exc()}")
                    else:  # Gemma 3
                        st.subheader("Результат перевода (Gemma 3):")
                        # Create a translator instance to check if text needs chunking
                        from utils.gemma_translation import GemmaTranslator
                        
                        try:
                            # Determine if we should use GPU based on user preference and permissions
                            gpu_layers = MODEL_CONFIG["gemma_gpu_layers"] if getattr(st.session_state, 'use_gpu', False) else 0
                            
                            # Create a new translator instance for this session
                            translator = GemmaTranslator()  # Each session gets its own model instance
                            if not translator.initialized:
                                # Use custom settings from environment
                                translator.load_model(
                                    n_gpu_layers=gpu_layers,
                                    context_size=MODEL_CONFIG["gemma_context_size"]
                                )
                            needs_chunking = translator.is_text_too_large(input_text)
                            
                            # Display appropriate spinner message based on whether chunking is needed
                            gpu_mode = "GPU" if translator.using_gpu else "CPU"
                            spinner_message = f"Processing text in chunks ({gpu_mode} mode)..." if needs_chunking else f"Processing translation ({gpu_mode} mode)..."
                            
                            with st.spinner(spinner_message):
                                try:
                                    # Instead of using the global display_streaming_translation function,
                                    # use a direct approach that keeps the translator instance
                                    result = ""
                                    translation_placeholder = st.empty()
                                    
                                    # Stream translation tokens and update UI
                                    for token in translator.translate_streaming(
                                        input_text, src_lang, tgt_lang,
                                        temperature=0.1, top_p=0.95
                                    ):
                                        result += token
                                        translation_placeholder.markdown(result)
                                        
                                except Exception as e:
                                    st.error(f"Translation error: {str(e)}")
                                    logging.error(f"Gemma translation error: {traceback.format_exc()}")
                                    result = None
                                
                                if result:
                                    # Prepare download capability
                                    doc = Document()
                                    doc.add_paragraph(result)
                                    doc_io = io.BytesIO()
                                    doc.save(doc_io)
                                    doc_io.seek(0)

                                    st.download_button(
                                        label="Скачать переведённый текст (.docx)",
                                        data=doc_io,
                                        file_name="gemma_translated_text.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                                else:
                                    st.warning("Не удалось выполнить перевод с Gemma 3.")
                                    
                                # Always unload model after use for better multi-user experience
                                translator.unload_model()
                        except Exception as model_error:
                            st.error(f"Model error: {str(model_error)}")
                            logging.error(f"Model error details: {traceback.format_exc()}")
                finally:
                    # Release the semaphore
                    model_semaphore.release()
            except Exception as outer_error:
                st.error(f"Unexpected error: {str(outer_error)}")
                logging.error(f"Unexpected error: {traceback.format_exc()}")
            finally:
                # Release translation lock
                st.session_state.translation_lock = False

def handle_readability_analysis():
    st.header("Анализ удобочитаемости текста")
    input_method = st.radio("Способ ввода текста:", ["Загрузить файл", "Вставить текст"])
    text = ""

    if input_method == "Загрузить файл":
        uploaded_file = st.file_uploader("Выберите файл (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"])
        if uploaded_file is not None:
            suffix = os.path.splitext(uploaded_file.name)[1]
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(uploaded_file.getbuffer())
                temp_file_path = tmp_file.name
            text = read_file(temp_file_path)
            os.remove(temp_file_path)
            st.write("**Содержимое файла:**")
            st.write(text)
    else:
        text = st.text_area("Вставьте ваш текст здесь", height=200)

    if text:
        auto_detect = st.checkbox("Определить язык автоматически", value=True)
        if auto_detect:
            detected_lang = detect_language(text)
            st.info(f"Определён язык: {detected_lang}")
            lang_code = detected_lang if detected_lang in ['ru','en','kk'] else 'en'
        else:
            lang_code = st.selectbox("Язык текста", ["ru", "en", "kk"])

        if st.button("Анализировать"):
            # Prevent multiple concurrent analyses
            if 'analysis_lock' in st.session_state and st.session_state.analysis_lock:
                st.warning("Анализ уже выполняется. Пожалуйста, дождитесь завершения.")
                return
                
            # Set analysis lock
            st.session_state.analysis_lock = True
            
            try:
                # Use the model semaphore for consistency with translation
                acquired = model_semaphore.acquire(blocking=False)
                if not acquired:
                    st.warning("Система загружена. Пожалуйста, попробуйте позже.")
                    st.session_state.analysis_lock = False
                    return
                    
                try:
                    with st.spinner("Выполняется анализ..."):
                        fre = flesch_reading_ease(text, lang_code)
                        fkgl = flesch_kincaid_grade_level(text, lang_code)
                        fog = gunning_fog_index(text, lang_code)
                        smog = smog_index(text, lang_code)

                    st.subheader("Результаты удобочитаемости")
                    st.markdown(
                        f"**Индекс удобочитаемости Флеша:** {color_code_index('Flesch Reading Ease', fre)}",
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        f"**Индекс Флеша-Кинкейда:** {color_code_index('Flesch-Kincaid Grade Level', fkgl)}",
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        f"**Индекс тумана Ганнинга:** {color_code_index('Gunning Fog Index', fog)}",
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        f"**Индекс SMOG:** {color_code_index('SMOG Index', smog)}",
                        unsafe_allow_html=True
                    )
                finally:
                    # Release the semaphore
                    model_semaphore.release()
            finally:
                # Release analysis lock
                st.session_state.analysis_lock = False

def main():
    setup_nltk()
    
    # Log the model configuration only once per session
    if 'model_config_logged' not in st.session_state:
        logging.info(f"Using model configuration: {MODEL_CONFIG}")
        st.session_state.model_config_logged = True

    st.title("Translation & Readability Analysis")
    st.sidebar.header("Функциональность")
    functionality = st.sidebar.radio("Выберите режим:", ["Перевод", "Анализ удобочитаемости"])

    if functionality == "Перевод":
        handle_translation()
    elif functionality == "Анализ удобочитаемости":
        handle_readability_analysis()

if __name__ == "__main__":
    main()